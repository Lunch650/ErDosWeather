#! python3
# coding=utf8
import requests
from datetime import datetime, timedelta
from docx import Document
from bs4 import BeautifulSoup


def pagecontent(code):
    # 获取页面数据
    url = 'http://wisdom.tqonline.top/weiqixiang/tianqi/getforecastbystations?stations=[' + code + ']'
    try:
        r = requests.get(url, timeout=10)
        weathersoup = BeautifulSoup(r.content, "html.parser")
        return weathersoup.body
    except requests.Timeout:
        print('呀，网络好像出现了一些问题，重新运行一次吧')
        return None


def pagelines(code):
    # 去除页面中不相干行数，返回干净的List数据
    templines = str(pagecontent(code)).split("<br/>")
    return [n.strip() for n in templines if (n.strip() != '') and (n.find('<') < 0)]


def weatherlines(prelist):
    # 根据页面的内容，判断发布的信息条数
    # 若没有信息则报错
    # 若有一条信息则返回该list
    # 若有两条信息则返回包含两条list的list
    if prelist.count('报歉，暂无预报数据。') == 2:
        print('报歉，' + str(prelist[0])[:-1] + '暂无预报数据。')
        return []
    elif prelist.count('报歉，暂无预报数据。') == 1:
        weatherweek = [[d[d.index('：') + 1:] for d in prelist[1:prelist[1:].index(prelist[0]) + 1]], ]
        return weatherweek
    else:
        weatherweek = \
            [
                [
                    d[d.index('：') + 1:] for d in prelist[1:prelist[1:].index(prelist[0]) + 1]
                ],
                [
                    s[s.index('：') + 1:] for s in prelist[prelist[1:].index(prelist[0]) + 2:]
                ],
            ]
        return weatherweek


def weatherdocx(weatherweek, docname, startrow):
    doc = Document(docname)
    t = doc.tables[0]
    for column, weatherday in enumerate(weatherweek):
        templist = weatherday.split('，')
        temperature = templist[2][4:-1] + '~' + templist[3][4:-2]
        weather = [templist[0], temperature, templist[1]]
        for row, elementinday in enumerate(weather):
            t.cell((startrow * 3) + row + 1, column + 3).text = elementinday
    doc.save(docname)


def predocx():
    # 由模板生成两个word文件，分别对应8点钟版本及16点版本，并修改模板中的时间
    d = Document('weatherTemplate.docx')
    d.paragraphs[1].text = '(' + dt.strftime('%Y') + '年第' + dt.strftime('%W') + '期）'
    d.paragraphs[3].text = dt.strftime('%Y') + '年' + dt.strftime('%m') + '月' + dt.strftime('%d') + '日'
    t = d.tables[0]
    for i in range(1, 8):
        dttemp = dt + timedelta(days=i)
        t.cell(0, i + 2).text = (dttemp.strftime('%m') + '月' + dttemp.strftime('%d') + '日')
    d.save(dt.strftime('%Y%m%d') + '08.docx')
    d.save(dt.strftime('%Y%m%d') + '16.docx')


if __name__ == '__main__':
    # citydict = {
    #     '准格尔旗': '53553',
    #     '清水河县': '53562',
    #     '和林格尔县': '53469',
    #     '凉城': '53475',
    #     '丰镇': '53484',
    #     '大同': '53487',
    #     '秦皇岛': '54449',
    #     '和林格尔县': '53469',
    #     '凉城': '53475',
    #     '右玉': '53478',
    #     '平鲁': '53574',
    #     '朔州': '53578',
    #     '神池': '53575',
    # }
    dt = datetime.now()
    citylist = \
        [
            '53553', '53562', '53469', '53475', '53484', '53487', '54449',
            '53469', '53475', '53478', '53574', '53578', '53575',
        ]
    predocx()
    for index, citycode in enumerate(citylist):
        weatherweeks = weatherlines(pagelines(citycode))
        if len(weatherweeks) > 0:
            # 如果页面有数据则先做出8点版本
            doc8name = dt.strftime('%Y%m%d') + '08.docx'
            weatherdocx(weatherweeks[0], doc8name, index)
            if len(weatherweeks) > 1:
                # 如果页面有8点及16点版本则更新
                d16name = dt.strftime('%Y%m%d') + '16.docx'
                weatherdocx(weatherweeks[1], d16name, index)