#!python3
# coding=utf8
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime, timedelta


class City(object):
    url = 'http://wisdom.tqonline.top/weiqixiang/tianqi/getforecastbystations'
    doc_names = [datetime.now().strftime('%Y%m%d') + '08.docx',
                 datetime.now().strftime('%Y%m%d') + '20.docx'
                 ]

    def __init__(self, city_name, city_code):
        self.__name = city_name
        self.__code = city_code
        self.__param = {'stations': '['+city_code+']'}

    def print_city_info(self):
        print('城市名称为:' + self.__name + ';城市代码为:' + self.__code)

    def content_from_page(self):
        # 利用request模块获取页面数据
        # 返回bs4模块解析后的body内容
        r = None
        while r is None:
            try:
                r = requests.get(City.url, params=self.__param, timeout=10)
            except requests.Timeout:
                print(self.__name + ':页面加载连接超时，加载任务重新启动中')
        return BeautifulSoup(r.content, 'html.parser').body

    def weekly_weather(mess_str):
        # 清洗页面返回的字符串内容,并根据内容将返回内容分开为两个list，返回内容格式为list
        if mess_str is not '':
            temp_split = str(mess_str).split("<br/>")
            temp_list = [n.strip() for n in temp_split if (n.strip() != '') and (n.find('<') < 0)]
            return City.separated_list(temp_list)
        else:
            print('错误')
            return None

    @staticmethod
    def separated_list(prelist):
        if prelist is []:
            print('找不到需要处理的列表，建议检查网站是否发布数据')
            return None
        if prelist.count('报歉，暂无预报数据。') == 2:
            print('报歉，' + str(prelist[0])[:-1] + '还未发布数据。')
            weekly_weather = []
        elif prelist.count('报歉，暂无预报数据。') == 1:
            weekly_weather = \
                [
                    [
                        d[d.index('：') + 1:] for d in prelist[1:prelist[1:].index(prelist[0]) + 1]
                    ],
                ]
        else:
            weekly_weather = \
                [
                    [
                        d[d.index('：') + 1:] for d in prelist[1:prelist[1:].index(prelist[0]) + 1]
                    ],
                    [
                        s[s.index('：') + 1:] for s in prelist[prelist[1:].index(prelist[0]) + 2:]
                    ],
                ]
        return weekly_weather

    @staticmethod
    def save_doc(doc_name, weekly_weather, start_row, path='', ):
        doc = Document(path + doc_name)
        t = doc.tables[0]
        for column, daily_weather in enumerate(weekly_weather):
            daily_weather = City.format_daily(daily_weather)
            for row, daily_element in enumerate(daily_weather):
                t.cell((start_row * 3) + row + 1, column + 3).text = daily_element
        doc.save(path + doc_name)

    @staticmethod
    def format_daily(daily_weather):
        templist = daily_weather.split('，')
        temperature = templist[2][4:-1] + '~' + templist[3][4:-2]
        return [templist[0], temperature, templist[1]]

    @staticmethod
    def predocx():
        # 由模板生成两个word文件，分别对应8点钟版本及20点版本，并修改模板中的时间
        d = Document('weatherTemplate.docx')
        dt = datetime.now()
        d.paragraphs[1].text = '(' + dt.strftime('%Y') + '年第' + dt.strftime('%W') + '期）'
        d.paragraphs[3].text = dt.strftime('%Y') + '年' + dt.strftime('%m') + '月' + dt.strftime('%d') + '日'
        t = d.tables[0]
        for column in range(1, 8):
            dttemp = dt + timedelta(days=column)
            t.cell(0, column + 2).text = (dttemp.strftime('%m') + '月' + dttemp.strftime('%d') + '日')
        for doc_name in City.doc_names:
            d.save(doc_name)


if __name__ == '__main__':
    def __main():
        cities = {
            '53553': '准格尔旗',
            '53562': '清水河县',
            '53484': '丰镇',
            '53487': '大同',
            '54449': '秦皇岛',
            '53469': '和林格尔县',
            '53475': '凉城',
            '53478': '右玉',
            '53574': '平鲁',
            '53578': '朔州',
            '53575': '神池',
        }
        cities_order = \
            [
                '53553', '53562', '53469', '53475', '53484', '53487', '54449',
                '53469', '53475', '53478', '53574', '53578', '53575',
            ]

        City.predocx()
        print('今日数据模板准备完毕')
        for i, code in enumerate(cities_order):
            city = City(cities.get(code), code)
            weather = City.weekly_weather(city.content_from_page())
            if len(weather) > 0:
                City.save_doc(doc_name=City.doc_names[0], weekly_weather=weather[0], start_row=i)
                print(cities.get(code), '8点版本写入完成')
                if len(weather) > 1:
                    City.save_doc(doc_name=City.doc_names[1], weekly_weather=weather[1], start_row=i)
                    print(cities.get(code), '20点版本写入完成')

    __main()
